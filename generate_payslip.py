"""
generate_payslip.py

Flow:
  1. payslip_template.docx is the master template — contains {{PLACEHOLDERS}}
     for all variable fields. It is created once by build_template_docx() if
     it doesn't already exist. You can edit its styling in Word freely.
  2. build_docx() loads the template, clones it, replaces every placeholder
     with real data, removes unused rows, and saves to output_path.
"""

import os
import re
import shutil
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(_BASE_DIR, 'payslip_template.docx')
LOGO_PATH     = os.path.join(_BASE_DIR, 'logo.jpg')
OUTPUT_DIR    = os.path.join(_BASE_DIR, 'output')
os.makedirs(OUTPUT_DIR, exist_ok=True)

COMPANY_NAME    = 'OTOPIA Batam 2'
COMPANY_ADDRESS = ('Pertokoan Kopkarlak Batam Centre, Jl. Raja Isa no 3A, '
                   'Baloi Permai, Kec. Batam Kota, Kota Batam, '
                   'Kepulauan Riau 29433, Indonesia')


# ── Formatting helpers ─────────────────────────────────────────────────────

def fmt_rp(amount):
    amount = int(round(float(amount or 0)))
    return 'Rp' + '{:,}'.format(amount).replace(',', '.')


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val)
            el.set(qn('w:sz'),    '4' if val == 'single' else '0')
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_spacing(para, before=0, after=0, line=240):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'),   str(before))
    spacing.set(qn('w:after'),    str(after))
    spacing.set(qn('w:line'),     str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def _set_font(run, name, size_pt, bold=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    name)
    rFonts.set(qn('w:hAnsi'),    name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'),       name)
    rPr.insert(0, rFonts)


def _cell_para(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT,
               font='Liberation Serif', size=12, bold=False):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    _para_spacing(para, before=0, after=0, line=240)
    run = para.add_run(text)
    _set_font(run, font, size, bold)
    return run


# ── Header ─────────────────────────────────────────────────────────────────

def _build_header(doc):
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p1, 0, 0, 240)
    if os.path.isfile(LOGO_PATH):
        p1.add_run().add_picture(LOGO_PATH, height=Cm(0.9))
    r1 = p1.add_run('  ' + COMPANY_NAME)
    _set_font(r1, 'Arial Black', 20)

    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p2, 0, 0, 240)
    _set_font(p2.add_run(COMPANY_ADDRESS), 'Arial', 9)

    header.add_paragraph()  # blank line

    p4 = header.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p4, 0, 0, 240)
    _set_font(p4.add_run('PAYSLIP'), 'Arial Black', 15)


# ── Info table ─────────────────────────────────────────────────────────────

def _build_info_table(doc, nama, jabatan, periode):
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(2.3)
    tbl.columns[1].width = Cm(14.5)
    for i, (label, value) in enumerate([('Nama', nama),
                                         ('Jabatan', jabatan),
                                         ('Periode', periode)]):
        is_first = (i == 0)
        for cell, text in zip(tbl.rows[i].cells, [label, value]):
            _cell_para(cell, text, font='Liberation Serif', size=12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell,
                             top='single' if is_first else 'nil',
                             bottom='single', left='nil', right='nil')


# ── Earnings table ─────────────────────────────────────────────────────────

def _build_earnings_table(doc, rows):
    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.style = 'Table Grid'
    for i, w in enumerate([Cm(6.48), Cm(3.58), Cm(4.01), Cm(2.97)]):
        tbl.columns[i].width = w
    for r_idx, (label, gaji, potongan, total, is_hdr) in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx, text in enumerate([label, gaji, potongan, total]):
            cell  = row.cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.LEFT if (c_idx == 0 or is_hdr) else WD_ALIGN_PARAGRAPH.RIGHT
            _cell_para(cell, text, align=align,
                       font='Liberation Serif' if (c_idx == 0 or is_hdr) else 'Arial',
                       size=12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell,
                             top='single' if is_hdr else 'nil',
                             bottom='single', left='nil', right='nil')


# ── Template creator ───────────────────────────────────────────────────────

def build_template_docx(output_path=None):
    """
    Create payslip_template.docx with {{PLACEHOLDER}} fields.
    This file defines the layout and styling. Edit it in Word to customise.
    """
    if output_path is None:
        output_path = TEMPLATE_PATH

    doc     = Document()
    section = doc.sections[0]
    section.page_height      = Cm(29.7)
    section.page_width       = Cm(21.0)
    section.top_margin       = Cm(6.2)
    section.bottom_margin    = Cm(2.0)
    section.left_margin      = Cm(2.0)
    section.right_margin     = Cm(2.0)
    section.header_distance  = Cm(1.0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

    _build_header(doc)

    p = doc.add_paragraph(); _para_spacing(p, 0, 0, 240)
    _build_info_table(doc, '{{NAMA}}', '{{JABATAN}}', '{{PERIODE}}')

    for _ in range(2):
        p = doc.add_paragraph(); _para_spacing(p, 0, 0, 240)

    _build_earnings_table(doc, [
        ('Penghasilan',              'Gaji',           'Potongan',   'Total',    True),
        ('Gaji Pokok',               '{{GAJI_POKOK}}', '',           '',         False),
        ('Komisi Coating',           '{{K_COATING}}',  '',           '',         False),
        ('Komisi Detailing',         '{{K_DETAILING}}','',           '',         False),
        ('Komisi Maintenance',       '{{K_MAINT}}',    '',           '',         False),
        ('Uang Makan',               '{{UANG_MAKAN}}', '',           '',         False),
        ('THR',                      '{{THR}}',        '',           '',         False),
        ('Cicilan Kasbon',           '',               '{{KASBON}}', '',         False),
        ('{{SUMMARY_LABEL}}',        '',               '',           '{{TOTAL}}',False),
    ])

    doc.save(output_path)
    return output_path


# ── Placeholder fill-in ────────────────────────────────────────────────────

def _replace_in_para(para, replacements):
    """Replace {{KEY}} tokens in a paragraph, preserving run formatting."""
    for run in para.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)


def _replace_in_doc(doc, replacements):
    """Walk all paragraphs and table cells and apply replacements."""
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    # Header paragraphs
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)


def _delete_table_row(table, row_idx):
    """Remove a row from a table by index."""
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)


def _remove_rows_with_placeholder(table, placeholder):
    """Delete any row whose cells still contain an unfilled placeholder."""
    rows_to_delete = []
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                if placeholder in para.text:
                    rows_to_delete.append(i)
                    break
    # Delete in reverse so indices stay valid
    for i in reversed(rows_to_delete):
        _delete_table_row(table, i)


# ── Main builder ───────────────────────────────────────────────────────────

def build_docx(emp, period, num_coating, num_detailing, num_maintenance,
               uang_makan_days, thr, kasbon, output_path):

    # Ensure template exists
    if not os.path.isfile(TEMPLATE_PATH):
        build_template_docx()

    # ── compute ──
    gaji_pokok    = float(emp['Gaji Pokok'] or 0)
    rate_coat     = float(emp['komisi coating detailing'] or 0)
    rate_detail   = float(emp['komisi detailing only'] or 0)
    rate_maint    = float(emp['komisi maintenance'] or 0)
    rate_um       = float(emp['uang makan 1 hari'] or 0)

    komisi_coat   = num_coating    * rate_coat
    komisi_detail = num_detailing  * rate_detail
    komisi_maint  = num_maintenance * rate_maint
    uang_makan    = uang_makan_days * rate_um
    subtotal      = gaji_pokok + komisi_coat + komisi_detail + komisi_maint + uang_makan + float(thr or 0)
    net           = subtotal - float(kasbon or 0)

    summary_label = ('Gaji+Komisi+Uang Makan - Kasbon'
                     if uang_makan_days > 0 else 'Gaji+Komisi - Kasbon')

    # ── load template as fresh copy ──
    doc = Document(TEMPLATE_PATH)

    # ── build replacements dict ──
    replacements = {
        '{{NAMA}}':         emp['Nama'],
        '{{JABATAN}}':      emp.get('Jabatan') or '-',
        '{{PERIODE}}':      period,
        '{{GAJI_POKOK}}':   fmt_rp(gaji_pokok),
        '{{K_COATING}}':    fmt_rp(komisi_coat)   if num_coating    > 0 else '',
        '{{K_DETAILING}}':  fmt_rp(komisi_detail) if num_detailing  > 0 else '',
        '{{K_MAINT}}':      fmt_rp(komisi_maint)  if num_maintenance > 0 else '',
        '{{UANG_MAKAN}}':   fmt_rp(uang_makan)    if uang_makan_days > 0 else '',
        '{{THR}}':          fmt_rp(thr)            if float(thr or 0) > 0 else '',
        '{{KASBON}}':       fmt_rp(kasbon)         if float(kasbon or 0) > 0 else '',
        '{{TOTAL}}':        fmt_rp(net),
        '{{SUMMARY_LABEL}}': summary_label,
    }

    # ── fill placeholders ──
    _replace_in_doc(doc, replacements)

    # ── remove rows whose value was set to '' (unused optional rows) ──
    # Any cell still containing an empty placeholder means the row was zeroed out
    # We detect them by checking for rows where the value cell is blank after fill
    optional_placeholders = [
        '{{K_COATING}}', '{{K_DETAILING}}', '{{K_MAINT}}',
        '{{UANG_MAKAN}}', '{{THR}}', '{{KASBON}}',
    ]
    for table in doc.tables:
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            row_text = ' '.join(
                para.text for cell in row.cells for para in cell.paragraphs
            )
            # If a row's combined text is just the label with no value
            # and all value cells are empty, drop it
            all_values_empty = all(
                cell.paragraphs[0].text.strip() == ''
                for cell in row.cells[1:]
            )
            is_optional = any(
                ph.replace('{{','').replace('}}','') in row_text
                for ph in optional_placeholders
            )
            # A row is removable if value columns are all empty AND
            # it's one of our optional rows (not header, not Gaji Pokok, not Total)
            label_text = row.cells[0].paragraphs[0].text.strip()
            removable_labels = {
                'Komisi Coating', 'Komisi Detailing', 'Komisi Maintenance',
                'Uang Makan', 'THR', 'Cicilan Kasbon',
            }
            if label_text in removable_labels and all_values_empty:
                rows_to_delete.append(i)
        for i in reversed(rows_to_delete):
            _delete_table_row(table, i)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    # Generate the template first
    t = build_template_docx()
    print('Template created:', t)

    # Test with real data
    from app import load_employees
    emps = load_employees()
    for emp in emps:
        safe = emp['Nama'].replace(' ', '_')
        path = os.path.join(OUTPUT_DIR, f'Test_{safe}.docx')
        build_docx(emp, 'Mei 2026', 5, 3, 2, 26, 0, 2650000 if emp['Nama'] == 'Agung Pratama' else 0, path)
        print('Generated:', path)
